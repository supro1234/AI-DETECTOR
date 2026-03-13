import sys
import json
import os
import base64
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_report(data, output_path):
    doc = Document()

    # --- Header ---
    header_section = doc.sections[0]
    header = header_section.header
    p = header.paragraphs[0]
    p.text = "AI IMAGE DETECTOR - FORENSIC REPORT"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Title ---
    title = doc.add_heading('Forensic Investigation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Summary Table ---
    doc.add_heading('Case Summary', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    summary_data = [
        ("File Name", data.get("file_name", "N/A")),
        ("Analysis ID", data.get("id", "N/A")),
        ("Timestamp", data.get("analyzed_at", datetime.now().isoformat())),
        ("Verdict", data.get("verdict", "Unknown").upper()),
        ("Veracity Score", f"{100 - data.get('confidence_score', 0)}%")
    ]

    for label, value in summary_data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        # Apply bold to verdict if dangerous
        if label == "Verdict" and ("FAKE" in value or "SUSPICIOUS" in value or "GENERATED" in value):
             run = row_cells[1].paragraphs[0].runs[0]
             run.font.color.rgb = RGBColor(255, 0, 0)
             run.bold = True

    # --- Forensic Details ---
    doc.add_heading('Neural Analysis Results', level=1)
    
    mb = data.get("model_breakdown", {})
    doc.add_paragraph(f"Neural Pixel Radix (NPR): {100 - mb.get('npr', 0)}%")
    doc.add_paragraph(f"Unique Frequency Detachment (UFD): {100 - mb.get('ufd', 0)}%")
    doc.add_paragraph(f"Cross-Efficient ViT: {100 - mb.get('crossvit', 0)}%")

    # --- Explanation ---
    doc.add_heading('Forensic Synthesis', level=1)
    doc.add_paragraph(data.get("explanation", "No detailed explanation provided."))

    # --- Forensic Points ---
    fp = data.get("forensic_points", {})
    if fp:
        doc.add_heading('Parameter Checklist', level=2)
        for key, value in fp.items():
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(f"{key.replace('_', ' ').title()}: ")
            run.bold = True
            p.add_run(value)

    # --- Red Flags & Authentic Signals ---
    flags = data.get("key_red_flags", [])
    if flags:
        doc.add_heading('Key Red Flags', level=2)
        for flag in flags:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(flag)
            run.font.color.rgb = RGBColor(200, 0, 0)

    signals = data.get("key_authentic_signals", [])
    if signals:
        doc.add_heading('Authentic Signals', level=2)
        for signal in signals:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(signal)
            run.font.color.rgb = RGBColor(0, 150, 0)

    # --- Evidence Image ---
    image_url = data.get("imageUrl", "")
    if image_url:
        try:
            doc.add_heading('Evidence Image', level=1)
            img_bytes = None

            if image_url.startswith("data:"):
                # Base64 data URL: data:image/jpeg;base64,....
                header, encoded = image_url.split(",", 1)
                img_bytes = base64.b64decode(encoded)
            elif image_url.startswith("http"):
                import urllib.request
                with urllib.request.urlopen(image_url, timeout=10) as r:
                    img_bytes = r.read()

            if img_bytes:
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(4.5))
                p = doc.paragraphs[-1]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = doc.add_paragraph(f"File: {data.get('file_name', 'Evidence')}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as img_err:
            doc.add_paragraph(f"[Image could not be embedded: {str(img_err)}]")

    # --- Footer ---
    doc.add_paragraph("\n\nProduced by AI Image Detector Forensic Suite.")
    
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python report_generator.py <output_path>  (JSON data via stdin)")
        sys.exit(1)

    try:
        output_path = sys.argv[1]
        
        # Read full JSON from stdin (piped from Node.js - no size limits)
        raw = sys.stdin.read()
        data = json.loads(raw)
        
        create_report(data, output_path)
        print(f"SUCCESS:{output_path}")
    except Exception as e:
        import traceback
        print(f"ERROR:{str(e)}\n{traceback.format_exc()}")
        sys.exit(1)
